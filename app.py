from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import io
import base64
import zipfile
import logging

# Logging setup (Error দেখার জন্য)
logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)

# --- CORS FIX (সব অরিজিন অ্যালাউ করার জন্য) ---
CORS(app, resources={r"/*": {"origins": "*"}})

def decode_image(data_url):
    """Base64 string থেকে ইমেজ বাইনারি ডেটায় কনভার্ট করে"""
    if not data_url:
        return None
    try:
        header, encoded = data_url.split(",", 1)
        return io.BytesIO(base64.b64decode(encoded))
    except Exception as e:
        print(f"Image decode error: {e}")
        return None

def create_set_document(set_data, set_name):
    doc = Document()
    
    heading = doc.add_heading(f'{set_name}', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for q in set_data['questions']:
        p = doc.add_paragraph()
        p.add_run(f"{q['questionNumber']}. {q['questionText']}").bold = True
        
        if q.get('questionImage'):
            image_stream = decode_image(q['questionImage'])
            if image_stream:
                try:
                    doc.add_picture(image_stream, width=Inches(3.0))
                except Exception as e:
                    print(f"Error adding question image: {e}")

        for opt in q['options']:
            p_opt = doc.add_paragraph(style='List Bullet')
            p_opt.add_run(f"{opt['letter']}) {opt['text']}")
            
            if opt.get('image'):
                opt_img_stream = decode_image(opt['image'])
                if opt_img_stream:
                    try:
                        doc.add_picture(opt_img_stream, width=Inches(1.5))
                    except Exception as e:
                        print(f"Error adding option image: {e}")
        
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading(f'Answer Key - {set_name}', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question No'
    hdr_cells[1].text = 'Correct Option'

    for q in set_data['questions']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(q['questionNumber'])
        row_cells[1].text = str(q['correctAnswer'])

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

# --- HOME ROUTE (Render চেক করার জন্য) ---
@app.route('/', methods=['GET'])
def home():
    return "MCQ Bot Backend is Live! Ready to generate sets."

# --- GENERATE SETS ROUTE (গুরুত্বপূর্ণ) ---
@app.route('/generate-sets', methods=['POST', 'OPTIONS'])
def generate_sets():
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'}), 200
        
    try:
        print("Request received for generate-sets") # Logs
        data = request.json
        
        if not data:
            return jsonify({'error': 'No data received'}), 400

        questions = data.get('questions', [])
        num_sets = data.get('numSets', 1)
        
        memory_file = io.BytesIO()
        
        with zipfile.ZipFile(memory_file, 'w') as zf:
            for i in range(num_sets):
                set_name = f"Set {chr(65 + i)}"
                
                # Logic to prevent crash if questions are empty
                if not questions:
                    continue

                shuffled_qs = random.sample(questions, len(questions))
                processed_questions = []
                
                for idx, q in enumerate(shuffled_qs):
                    original_options = q['options']
                    # Safety check for index
                    correct_idx = int(q['correctAnswer']) - 1
                    if correct_idx < 0 or correct_idx >= len(original_options):
                        correct_idx = 0 # Fallback
                        
                    correct_opt_id = original_options[correct_idx]['id']
                    
                    shuffled_opts = random.sample(original_options, len(original_options))
                    new_correct_idx = -1
                    formatted_opts = []
                    
                    for opt_idx, opt in enumerate(shuffled_opts):
                        if opt['id'] == correct_opt_id:
                            new_correct_idx = opt_idx
                        
                        formatted_opts.append({
                            'letter': chr(97 + opt_idx),
                            'text': opt['text'],
                            'image': opt['image']
                        })
                    
                    processed_questions.append({
                        'questionNumber': idx + 1,
                        'questionText': q['questionText'],
                        'questionImage': q['questionImage'],
                        'options': formatted_opts,
                        'correctAnswer': chr(97 + new_correct_idx)
                    })

                set_data = {
                    'setName': set_name,
                    'questions': processed_questions
                }
                
                docx_file = create_set_document(set_data, set_name)
                zf.writestr(f"{set_name}.docx", docx_file.getvalue())

        memory_file.seek(0)  
        
        return send_file(
            memory_file,
            download_name='mcq_sets.zip',
            as_attachment=True,
            mimetype='application/zip'
        )

    except Exception as e:
        print(f"CRITICAL ERROR: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Gunicorn প্রোডাকশনে এটি ব্যবহার করবে না, কিন্তু লোকাল রান করার জন্য ঠিক আছে
    app.run(debug=True, port=5000)